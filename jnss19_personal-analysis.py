import pandas as pd
from docx import Document
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
import pymysql
from sqlalchemy import create_engine
import re
import os
readPath=r'C:/Users/22828/Desktop/问卷得分情况/'
savePath=r'C:/Users/22828/Desktop/个人分析/'
fileList=os.listdir(readPath)
# col=data.columns.to_list()
# for i in range(len(col)):
#     print(i,col[i])
for fileName in fileList:
    print(fileName)
    data=pd.read_excel(readPath+fileName)
    document=Document()
    
    if '二级学科' not in data.columns and '培养方式' not in data.columns:
        i=0
    elif '二级学科' in data.columns and '培养方式' in data.columns:
        i=2
    elif '二级学科' not in data.columns and '培养方式' in data.columns:
        i=1
    elif '二级学科' in data.columns and '培养方式' not in data.columns:
        i=1
    else:
        print('error')
    print(i)
    head=document.add_heading('',level=1)
    run=head.add_run('一、单科风险情况梳理')
    run.font.name=u'宋体'
    run.font.color.rgb = RGBColor(0,0,0) #字体颜色
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.add_paragraph('')

    if data.iloc[[0],[6+i]].values[0][0]==0:
        document.add_paragraph('政治单科有风险，建议选择政治单科线较低的学校，并排除政治单科线较高的学校。')
    if data.iloc[[0],[5+i]].values[0][0]==0:
        document.add_paragraph('英语单科有风险，建议选择英语单科线较低的学校，并排除英语单科线较高的学校。')
    if data.iloc[[0],[7+i]].values[0][0]==0:
        document.add_paragraph('数学单科有风险，建议选择数学单科线较低的学校，并排除数学单科线较高的学校。')
    if data.iloc[[0],[8+i]].values[0][0]==0:
        document.add_paragraph('专业课有风险，建议选择专业课线较低的学校，并排除专业课线较高的学校。')
    if data.iloc[[0],[18+i]].values[0][0]==1:
        document.add_paragraph('复试专业课笔试有风险，建议选择复试专业课笔试线较低的学校，并排除复试专业课笔试线较高的学校。')
    if data.iloc[[0],[19+i]].values[0][0]==1:
        document.add_paragraph('复试面试有风险，建议选择复试面试线较低的学校，并排除复试面试线较高的学校。')
    if data.iloc[[0],[17+i]].values[0][0]==0:
        document.add_paragraph('复试英语有风险，建议选择复试英语线较低的学校，并排除复试英语线较高的学校。')
#     else:
#         document.add_paragraph('您的初试无风险')
        
    head=document.add_heading('',level=1)
    run=head.add_run('二、初试复试有劣势情况梳理')
    run.font.name=u'宋体'
    run.font.color.rgb = RGBColor(0,0,0) #字体颜色
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.add_paragraph('')

    #初试判断
    if data.iloc[[0],[20+i]].values[0][0]==1 or (data.iloc[[0],[20+i]].values[0][0]==0 and data.iloc[[0],[4+i]].values[0][0]==0):
        first_info='初试差'
    elif data.iloc[[0],[20+i]].values[0][0]==0 and (data.iloc[[0],[4+i]].values[0][0]==1 or data.iloc[[0],[4+i]].values[0][0]==2):
        first_info='初试一般'
    elif data.iloc[[0],[20+i]].values[0][0]==2 or (data.iloc[[0],[20+i]].values[0][0]==0 and data.iloc[[0],[4+i]].values[0][0]==5):
        first_info='初试好'
    else:
        print('error')

    #复试判断
    maxlist=[]
    for n in range(9+i,18+i):
        maxlist.append(data.iloc[[0],[n]].values[0][0])
    #26-34题的最大值
    max_26_34=max(maxlist)

    if data.iloc[[0],[20+i]].values[0][0]==2 or (data.iloc[[0],[20+i]].values[0][0]==0 and data.iloc[[0],[16+i]].values[0][0]==0 and max_26_34<4) or \
    max_26_34==0 or max_26_34==1:
        retest_info='复试差'
    elif (data.iloc[[0],[20+i]].values[0][0]==0 and (data.iloc[[0],[16+i]].values[0][0]==1 or data.iloc[[0],[16+i]].values[0][0]==2) and \
          (max_26_34==2 or max_26_34==3)) or (data.iloc[[0],[16+i]].values[0][0]==0 and max_26_34==4):
        retest_info='复试一般'
    elif data.iloc[[0],[20+i]].values[0][0]==1 or \
    (data.iloc[[0],[20+i]].values[0][0]==0 and (data.iloc[[0],[16+i]].values[0][0]==1 or data.iloc[[0],[16+i]].values[0][0]==2) and (max_26_34==4 or max_26_34==5))\
    or (data.iloc[[0],[16+i]].values[0][0]==0 and max_26_34==5):
        retest_info='复试好'
    else:
        print('error')
        
    if data.iloc[[0],[12+i]].values[0][0]==2:
        job1='班委成员'
    elif data.iloc[[0],[12+i]].values[0][0]==4:
        job1='班长'
    else:
        job1=''

    if data.iloc[[0],[13+i]].values[0][0]==2:
        job2='干事或副部长'
    elif data.iloc[[0],[13+i]].values[0][0]==4:
        job2='部长'
    elif data.iloc[[0],[13+i]].values[0][0]==5:
        job2='主席'
    else:
        job2=''

    if job2=='' and job1!='':
        job='积极担任过'+job1+'职务，'
    elif job2!='' and job1=='':
        job='积极担任过'+job2+'职务，'
    elif job2!='' and job1!='':
        job='积极担任过'+job1+'及'+job2+'职务，'
    else:
        job=''

    if data.iloc[[0],[9+i]].values[0][0]==1:
        match='参加过校级相关比赛，'
    elif data.iloc[[0],[9+i]].values[0][0]==3:
        match='参加过省级相关竞赛，'
    elif data.iloc[[0],[9+i]].values[0][0]==4:
        match='参加过国家级或以上相关比赛，'
    else:
        match=''

    if data.iloc[[0],[10+i]].values[0][0]==1 and data.iloc[[0],[9+i]].values[0][0]!=0:
         prize='且有获奖，'
    else:
        prize=''

    if data.iloc[[0],[11+i]].values[0][0]==5:
        publish='本科期间发表过学术论文或申请过发明专利，'
    else:
        publish=''

    if data.iloc[[0],[14+i]].values[0][0]==2:
        special='有某项特长并参加过校级及以上相关比赛，'
    elif data.iloc[[0],[14+i]].values[0][0]==4:
        special='参加过某些社团并担任骨干或组建过社团，'
    else:
        special=''

    if data.iloc[[0],[15+i]].values[0][0]==2:
        hobby='喜欢读书，涉猎广泛知识面广或对某领域颇有积累，'
    else:
        hobby=''

    if data.iloc[[0],[16+i]].values[0][0]==2:
        English_level='您已通过CET-6，'
    elif data.iloc[[0],[16+i]].values[0][0]==1:
        English_level='您已通过CET-4，'
    else:
        English_level=''

    if data.iloc[[0],[17+i]].values[0][0]==3:
        spoken_English='而且发音标准，交流轻松流畅，'
    elif data.iloc[[0],[17+i]].values[0][0]==2:
        spoken_English='而且能正常进行听说交流，'
    else:
        spoken_English=''
    
    if first_info=='初试差' and retest_info=='复试差':
        document.add_paragraph('您的初试和复试优势较为均衡，建议重点参考复试线和招生人数情况，选择分数线较低，招生人数多，尤其每年报考人数都不足，有调剂名额的学校。')
    elif first_info=='初试一般' and retest_info=='复试一般':
        document.add_paragraph('本科期间专业排名处于中上游，'+job+match+prize+publish+special+hobby+English_level+spoken_English+'初试和复试优势较为均衡，都有一定的实力。')
    elif first_info=='初试好' and retest_info=='复试好':
        document.add_paragraph('本科期间专业排名处于中上游，'+job+match+prize+publish+special+hobby+English_level+spoken_English+'初试和复试优势较为均衡，且都很突出，建议冲击较好学校。')
    elif first_info=='初试差' and retest_info=='复试一般':
        document.add_paragraph(job+match+prize+publish+special+hobby+English_level+spoken_English+'您的复试相比初试更有优势，建议选择复试成绩占比较高的学校专业。')
    elif first_info=='初试一般' and retest_info=='复试差':
        document.add_paragraph(job+match+prize+publish+special+hobby+English_level+spoken_English+'您的初试相比复试更有优势，建议选择初试成绩占比较高的学校专业。')
    elif first_info=='初试一般' and retest_info=='复试好':
        document.add_paragraph(job+match+prize+publish+special+hobby+English_level+spoken_English+'您的复试相比初试更有优势，建议选择复试成绩占比较高的学校专业。')
    elif first_info=='初试好' and retest_info=='复试一般':
        document.add_paragraph(job+match+prize+publish+special+hobby+English_level+spoken_English+'您的初试相比复试更有优势，建议选择初试成绩占比较高的学校专业。')
    elif first_info=='初试好' and retest_info=='复试差':
        document.add_paragraph(job+match+prize+publish+special+hobby+English_level+spoken_English+'您的初试相比复试更有优势，建议选择初试成绩占比较高的学校专业。')
    elif first_info=='初试差' and retest_info=='复试好':
        document.add_paragraph(job+match+prize+publish+special+hobby+English_level+spoken_English+'您的复试相比初试更有优势，建议选择复试成绩占比较高的学校专业。')
    else:
        print('error')
    
    head=document.add_heading('',level=1)
    run=head.add_run('三、复试笔试面试英语情况梳理')
    run.font.name=u'宋体'
    run.font.color.rgb = RGBColor(0,0,0) #字体颜色
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.add_paragraph('')

    #复试笔试判断
    if data.iloc[[0],[18+i]].values[0][0]==2 or \
    (data.iloc[[0],[18+i]].values[0][0]==0 and (data.iloc[[0],[4+i]].values[0][0]==5 or data.iloc[[0],[4+i]].values[0][0]==2)):
        written='笔试好'
    elif data.iloc[[0],[18+i]].values[0][0]==0 or data.iloc[[0],[4+i]].values[0][0]==1:
        written='笔试一般'
    elif data.iloc[[0],[18+i]].values[0][0]==1 or data.iloc[[0],[4+i]].values[0][0]==0:
        written='笔试差'
    else:
        print('error')

    maxlist=[]
    for n in range(10+i,16+i):
        maxlist.append(data.iloc[[0],[n]].values[0][0])
    #26-32题的最大值
    max_26_32=max(maxlist)

    #复试面试判断
    if data.iloc[[0],[19+i]].values[0][0]==2 or \
    (data.iloc[[0],[19+i]].values[0][0]==0 and (max_26_32==5 or max_26_32==4)):
        interview='面试好'
    elif data.iloc[[0],[19+i]].values[0][0]==0 or (data.iloc[[0],[19+i]].values[0][0]==0 and (max_26_32==2 or max_26_32==3)):
        interview='面试一般'
    elif data.iloc[[0],[19+i]].values[0][0]==1 or (data.iloc[[0],[19+i]].values[0][0]==0 and (max_26_32==0 or max_26_32==1)):
        interview='面试差'
    else:
        print('error')

    #复试英语判断
    if data.iloc[[0],[17+i]].values[0][0]==3:
        english='英语好'
    elif data.iloc[[0],[17+i]].values[0][0]==2 or \
    (data.iloc[[0],[17+i]].values[0][0]==1 and (data.iloc[[0],[16+i]].values[0][0]==1 or data.iloc[[0],[16+i]].values[0][0]==2)):
        english='英语一般'
    elif data.iloc[[0],[17+i]].values[0][0]==0 or (data.iloc[[0],[17+i]].values[0][0]==1 and data.iloc[[0],[16+i]].values[0][0]==0):
        english='英语差'
    else:
        print('error')
        
    if written=='笔试好' and interview=='面试差' and english=='英语差':
        document.add_paragraph('您的复试笔试实力较强，建议您选择复试笔试成绩占比较高的学校。')
    elif written=='笔试好' and interview=='面试一般' and english=='英语差':
        document.add_paragraph('您的复试笔试实力较强，建议您选择复试笔试成绩占比较高的学校。')
    elif written=='笔试好' and interview=='面试差' and english=='英语一般':
        document.add_paragraph('您的复试笔试实力较强，建议您选择复试笔试成绩占比较高的学校。')
    elif written=='笔试好' and interview=='面试一般' and english=='英语一般':
        document.add_paragraph('您的复试笔试实力较强，建议您选择复试笔试成绩占比较高的学校。')
    elif written=='笔试好' and interview=='面试好' and english=='英语差':
        document.add_paragraph('您的复试笔试、复试面试实力较强，建议您选择复试笔试成绩加复试面试成绩占比较高的学校。')
    elif written=='笔试好' and interview=='面试好' and english=='英语一般':
        document.add_paragraph('您的复试笔试、复试面试实力较强，建议您选择复试笔试成绩加复试面试成绩占比较高的学校。')
    elif written=='笔试好' and interview=='面试差' and english=='英语好':
        document.add_paragraph('您的复试笔试、英语实力较强，建议您选择复试笔试成绩加复试英语成绩占比较高的学校。')
    elif written=='笔试好' and interview=='面试一般' and english=='英语好':
        document.add_paragraph('您的复试笔试、英语实力较强，建议您选择复试笔试成绩加复试英语成绩占比较高的学校。')

    elif written=='笔试差' and interview=='面试好' and english=='英语差':
        document.add_paragraph('您的复试面试实力较强，建议您选择复试面试成绩占比较高的学校。')
    elif written=='笔试一般' and interview=='面试好' and english=='英语差':
        document.add_paragraph('您的复试面试实力较强，建议您选择复试面试成绩占比较高的学校。')
    elif written=='笔试差' and interview=='面试好' and english=='英语一般':
        document.add_paragraph('您的复试面试实力较强，建议您选择复试面试成绩占比较高的学校。')
    elif written=='笔试一般' and interview=='面试好' and english=='英语一般':
        document.add_paragraph('您的复试面试实力较强，建议您选择复试面试成绩占比较高的学校。')
    elif written=='笔试一般' and interview=='面试好' and english=='英语好':
        document.add_paragraph('您的复试面试、英语实力较强，建议您选择复试面试成绩加复试英语成绩占比较高的学校。')
    elif written=='笔试差' and interview=='面试好' and english=='英语好':
        document.add_paragraph('您的复试面试、英语实力较强，建议您选择复试面试成绩加复试英语成绩占比较高的学校。')

    elif written=='笔试差' and interview=='面试差' and english=='英语好':
        document.add_paragraph('您的复试英语实力较强，建议您选择复试英语成绩占比较高的学校。')
    elif written=='笔试差' and interview=='面试一般' and english=='英语好':
        document.add_paragraph('您的复试英语实力较强，建议您选择复试英语成绩占比较高的学校。')
    elif written=='笔试一般' and interview=='面试差' and english=='英语好':
        document.add_paragraph('您的复试英语实力较强，建议您选择复试英语成绩占比较高的学校。')
    elif written=='笔试一般' and interview=='面试一般' and english=='英语好':
        document.add_paragraph('您的复试英语实力较强，建议您选择复试英语成绩占比较高的学校。')

    elif written=='笔试一般' and interview=='面试一般' and english=='英语差':
        document.add_paragraph('您的复试笔试、复试面试实力较强，建议您选择复试笔试成绩加复试面试成绩占比较高的学校。')
    elif written=='笔试一般' and interview=='面试差' and english=='英语差':
        document.add_paragraph('您的复试笔试实力较强，建议您选择复试笔试成绩占比较高的学校。')
    elif written=='笔试一般' and interview=='面试差' and english=='英语一般':
        document.add_paragraph('您的复试笔试、英语实力较强，建议您选择复试面试成绩加复试英语成绩占比较高的学校。')    
    elif written=='笔试差' and interview=='面试一般' and english=='英语差':
        document.add_paragraph('您的复试面试实力较强，建议您选择复试面试成绩占比较高的学校。')
    elif written=='笔试差' and interview=='面试一般' and english=='英语一般':
        document.add_paragraph('您的复试面试、英语实力较强，建议您选择复试面试成绩加复试英语成绩占比较高的学校。')
    elif written=='笔试差' and interview=='面试差' and english=='英语一般':
        document.add_paragraph('您的复试英语实力较强，建议您选择复试英语成绩占比较高的学校。')

    elif written=='笔试差' and interview=='面试差' and english=='英语差':
        document.add_paragraph('您的复试笔试、复试面试、英语实力较均衡。')
    elif written=='笔试一般' and interview=='面试一般' and english=='英语一般':
        document.add_paragraph('您的复试笔试、复试面试、英语实力较均衡。')
    elif written=='笔试好' and interview=='面试好' and english=='英语好':
        document.add_paragraph('您的复试笔试、复试面试、英语实力较均衡。')
    else:
        print('error')
    
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(10.5) #字号
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0) #字体颜色
    document.save(savePath+fileName.replace('.xlsx','-个人分析')+'.docx')
