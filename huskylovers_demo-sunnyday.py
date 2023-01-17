# This Python 3 environment comes with many helpful analytics libraries installed

# It is defined by the kaggle/python docker image: https://github.com/kaggle/docker-python

# For example, here's several helpful packages to load in 



import numpy as np # linear algebra

import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)



# Input data files are available in the "../input/" directory.

# For example, running this (by clicking run or pressing Shift+Enter) will list the files in the input directory



import os

print(os.listdir("../input/new-wordsds/research_data/research_data/"))



# Any results you write to the current directory are saved as output.
!pip install python-docx
import docx

import re

from tqdm import tqdm

import jieba

import gensim

from collections import defaultdict 

from math import log

import re

import seaborn as sns

import matplotlib.pyplot as plt

positive_word = []

with open('../input/postive-word/positive_dict.txt','r') as fr:

    for line in fr.readlines():

        positive_word.append(line.replace('\n',''))

print(len(positive_word))

positive_word[:5]
word_len = [len(word) for word in positive_word]

sns.countplot(word_len)

plt.title('char_level length of the special word')

plt.show()
split_words = []

for word in positive_word:

    word = jieba.cut(word)

    try_data = ' '.join(word).split()

    split_words.append(try_data)
split_word_length = [len(word) for word in split_words]

sns.countplot(split_word_length)

plt.title('word_level length of the special word')

plt.show()
def read_data(file_path):

    text = []

    none = 0

    doc = docx.Document(file_path)

    #para

    for para in doc.paragraphs:

        content = para.text

        filter_ = re.compile(u'[^\u4E00-\u9FA5]')

        filtered_content = filter_.sub(r'',content)

        if len(filtered_content) > 0:

            text.append(filtered_content)

        else:

            none += 1

    #table

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                content = cell.text

                filter_ = re.compile(u'[^\u4E00-\u9FA5]')

                filtered_content = filter_.sub(r'',content)

                if len(filtered_content) > 0:

                    text.append(filtered_content)

                else:

                    none += 1

    #print(str(none)+'空行')

    return text
path = "../input/new-wordsds/research_data/research_data/"

doc = os.listdir("../input/new-wordsds/research_data/research_data/")

doc
text =[]

for doc_ in tqdm(doc):

    path_ = path + doc_

    text_ = read_data(path_)

    text += text_

    
text_word = []

words = set([])

for text_ in tqdm(text):

    text_word.append(list(jieba.cut(text_)))

    words = words | set(jieba.cut(text_))

words = list(words)
class Find_Words:

    def __init__(self, min_count=5, min_pmi=0):

        self.min_count = min_count

        self.min_pmi = min_pmi

        self.chars, self.pairs = defaultdict(int), defaultdict(int) #如果键不存在，那么就用int函数

                                                                  #初始化一个值，int()的默认结果为0

        self.total = 0.

    def count(self, texts): #计数函数，计算单字出现频数、相邻两字出现的频数

        for text in texts:

            self.chars[text[0]] += 1

            for i in range(len(text)-1):

                self.chars[text[i+1]] += 1

                self.pairs[text[i:i+2]] += 1

                self.total += 1

        self.chars = {i:j for i,j in self.chars.items() if j >= self.min_count} #最少频数过滤

        self.pairs = {i:j for i,j in self.pairs.items() if j >= self.min_count} #最少频数过滤

        self.strong_segments = set()

        for i,j in self.pairs.items(): #根据互信息找出比较“密切”的邻字

            _ = log(self.total*j/(self.chars[i[0]]*self.chars[i[1]]))

            if _ >= self.min_pmi:

                self.strong_segments.add(i)

    def find_words(self, texts): #根据前述结果来找词语

        self.words = defaultdict(int)

        for text in texts:

            s = text[0]

            for i in range(len(text)-1):

                if text[i:i+2] in self.strong_segments: #如果比较“密切”则不断开

                    s += text[i+1]

                else:

                    self.words[s] += 1 #否则断开，前述片段作为一个词来统计

                    s = text[i+1]

        self.words = {i:j for i,j in self.words.items() if j >= self.min_count and len(i)>2 and len(i)<7} #最后再次根据频数和字的数量

fw = Find_Words(5, 1)

fw.count(text)

fw.find_words(text)
new_words = set(fw.words)- set(words)

print(len(new_words))

list(new_words)[:100]
class Find_Words:

    def __init__(self, min_count=5, min_pmi=0):

        self.min_count = min_count

        self.min_pmi = min_pmi

        self.words, self.pairs, self.thr_pairs = defaultdict(int), defaultdict(int), defaultdict(int) #如果键不存在，那么就用int函数

                                                                  #初始化一个值，int()的默认结果为0

        self.total = 0.

        

    def count_2(self, texts): #计数函数，计算单字出现频数、相邻两字出现的频数

        for text in texts:

            self.words[text[0]] += 1

            for i in range(len(text)-1):

                self.words[''.join(text[i+1])] += 1

                self.pairs[' '.join(text[i:i+2])] += 1

                self.total += 1

        self.words = {i:j for i,j in self.words.items() if j >= self.min_count} #最少频数过滤

        self.pairs = {i:j for i,j in self.pairs.items() if j >= self.min_count} #最少频数过滤

        self.strong_segments = set()

        for i,j in self.pairs.items(): #根据互信息找出比较“密切”的邻字

            former,later = i.split()[0],i.split()[1]

            _ = log(self.total*j/(self.words[former]*self.words[later]))

            if _ >= self.min_pmi:

                self.strong_segments.add(i.replace(' ',''))

                

    def count_3(self, texts): #计数函数，计算单字出现频数、相邻两字出现的频数

        for text in texts:

            for i in range(len(text)-2):

                self.thr_pairs[' '.join(text[i:i+3])] += 1

                self.total += 1

        self.words = {i:j for i,j in self.words.items() if j >= self.min_count} #最少频数过滤

        self.thr_pairs = {i:j for i,j in self.thr_pairs.items() if j >= self.min_count} #最少频数过滤

        for i,j in self.thr_pairs.items(): #thr_pairs根据互信息找出比较“密切”的邻字

            former,later = ' '.join(i.split()[:2]),' '.join(i.split()[1:])

            former_word, later_word = i.split()[2], i.split()[0]

            try:

                try_former = log(self.total*j/(self.pairs[former]*self.words[former_word]))

                try_later = log(self.total*j/(self.pairs[later]*self.words[later_word]))

                _ = max([try_former,try_later])

                if _ >= self.min_pmi:

                    self.strong_segments.add(i.replace(' ',''))

            except:

                continue
fw = Find_Words(5, 1)

fw.count_2(text_word)

fw.count_3(text_word)
len(fw.strong_segments)
print(len(fw.strong_segments - set(words)))

new_words = new_words | set(fw.strong_segments - set(words))
split_sentence = []

for sentence in tqdm(text):

    split_sentence.append(' '.join(jieba.cut(sentence)).split())
two_combined_sentence = []

for sentence in tqdm(split_sentence):

    try_sentence = []

    for i in range(len(sentence)-1):

         try_sentence.append(''.join(sentence[i:i+2]))

        

    two_combined_sentence.append(try_sentence)

    
thr_combined_sentence = []

for sentence in tqdm(split_sentence):

    try_sentence = []

    for i in range(len(sentence)-2):

        try_sentence.append(''.join(sentence[i:i+3]))

    thr_combined_sentence.append(try_sentence)
four_combined_sentence = []

for sentence in tqdm(split_sentence):

    try_sentence = []

    for i in range(len(sentence)-3):

        try_sentence.append(''.join(sentence[i:i+4]))

    four_combined_sentence.append(try_sentence)
prepare_word = set([])

prepare_sentence = two_combined_sentence + thr_combined_sentence + four_combined_sentence

#prepare_sentence = split_sentence + two_combined_sentence + thr_combined_sentence + four_combined_sentence

for sentence in tqdm(prepare_sentence):

    prepare_word = prepare_word | set(sentence) 
i = 0

for word in new_words:

    if word in positive_word:

        i += 1

i
len(new_words)
i = 0

for word in prepare_word:

    if word in positive_word:

        i += 1

print(i)
len(prepare_word)
len(positive_word)